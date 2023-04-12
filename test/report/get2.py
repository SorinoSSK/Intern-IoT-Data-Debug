import requests
import time
import json
import pandas as pd
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import re
import os
import glob
import uuid
from docx import Document
from docx.shared import Inches, Mm
from docxtpl import DocxTemplate, InlineImage


class Config:
    def __init__(self, name="template.docx"):
        self.name = name
        self.units = self.read_template()

    readFlag = False

    def read_template(self):
        units = []
        document = Document(self.name)
        for p in document.paragraphs:
            # print(p.text)
            matches = re.findall("(?<=\{\{)(.*?)(?=\}\})", str(p.text))
            for match in matches:
                # print(match)
                match = match[1:-1]
                if 'unit' in match:
                    try:
                        [unit_id, unit_name] = match.split('_')[1:]
                        print(unit_id, unit_name)
                        units.append([unit_id, unit_name])
                    except ValueError:
                        print("unit number is not integer")
        self.readFlag = True
        return units

    def get_units(self):
        if self.readFlag:
            return self.units
        else:
            print("read template first")


class Unit():
    def __init__(self, id, name):
        self.id = id
        self.name = name
        self.token = self.get_token()
        self.df = None
        self.success = False

    def get_token(self):
        # get auth token
        url = "https://datakrewtech.com/api/sign-in"
        myobj = {'email': 'stevenedbert47@gmail.com', 'password': 'vflow123'}
        x = requests.post(url, data=myobj)

        token = x.json()['access_token']
        return token

    def get_data(self, duration):
        # get auth token
        # url = "https://datakrewtech.com/api/sign-in"
        # myobj = {'email': 'stevenedbert47@gmail.com', 'password': 'vflow123'}
        # x = requests.post(url, data=myobj)

        # token = x.json()['access_token']

        curr_time = int(time.time() * 1000)

        start_time = curr_time - 1000 * duration

        endpoint = 'https://datakrewtech.com/api/iot_mgmt/orgs/3/projects/70/gateways/' + \
            str(self.id) + '/data_dump_index'
        headers = {'Authorization': f'Bearer {self.token}'}
        params = {'page_size': 100000, 'page_number': 1,
                  'to_date': curr_time, 'from_date': start_time}

        response = requests.get(endpoint, headers=headers, params=params)
        try:
            json_dump = response.json()
        except json.JSONDecodeError:
            print('Key does not exist!')
            return

        if not json_dump['total_entries']:
            print('No data for this key!')
            return
        self.success = True
        df = pd.json_normalize(json_dump, 'data_dumps')
        unwanted = df.columns[df.columns.str.startswith('data.') == False]

        df.drop(unwanted, axis=1, inplace=True)
        df = df.rename(columns=lambda x: x.replace('data.assets_params.', ''))
        df = df.rename(columns=lambda x: x.replace('data.', ''))

        # invert dataframe so the most recent data is at the bottom
        df = df[:: -1]
        df = df.reset_index(drop=True)

        # convert timestamp to sg time
        df['timestamp'] = pd.to_datetime(df['timestamp'], unit='s', utc=True)
        df['timestamp'] = df['timestamp'].dt.tz_convert('Asia/Singapore')

        self.df = df

        if df is None:
            print("No data obtained, please try again")
        else:
            self.create_excel()

        return df

    def create_excel(self):
        os.makedirs("excel", exist_ok=True)
        filename = "excel/" + str(self.name) + ".xlsx"
        self.df['timestamp'] = self.df['timestamp'].dt.tz_localize(None)
        self.df.to_excel(filename)

    def return_columns(self):
        return list(self.df.columns)

    def plot(self, params, multiple_axes=False):
        if not self.success:
            print("No data obtained, ploting for unit " +
                  self.name + " is not possible")
            return
        os.makedirs("figures", exist_ok=True)
        if type(params) == str:
            print('Enter a list of parameters to plot')
            return
        if multiple_axes:
            fig = make_subplots(specs=[[{"secondary_y": True}]])
            fig.add_trace(
                go.Scatter(x=self.df['timestamp'], y=self.df[params[0]],
                           name=str(params[0])),
                secondary_y=False,
            )

            fig.add_trace(
                go.Scatter(x=self.df['timestamp'], y=self.df[params[1]],
                           name=str(params[1])),
                secondary_y=True,
            )

            fig.update_yaxes(title_text=str(params[0]), secondary_y=False)
            fig.update_yaxes(title_text=str(params[1]), secondary_y=True)

        else:
            fig = go.Figure()
            for param in params:
                fig.add_trace(go.Scatter(
                    x=self.df['timestamp'],
                    y=self.df[param],
                    name=param,
                    mode='lines'
                ))
                fig.update_layout(
                    yaxis_title=param,
                )

        fig.update_layout(
            height=1000,
            width=1250,
            showlegend=True,
            xaxis_linecolor='black',
            xaxis_title="Time",
            yaxis_linecolor='black',
            plot_bgcolor='white',
            font=dict(
                color="black",
                size=24
            )

        )

        filename = "figures/" + str(self.name) + \
            "_" + str(uuid.uuid4())[0:8] + ".png"
        fig.write_image(filename)


def get_seconds(str_time):
    if str_time[1].lower() == 's':
        return int(str_time[0])
    elif str_time[1].lower() == 'm':
        return int(str_time[0]) * 60
    elif str_time[1].lower() == 'h':
        return int(str_time[0]) * 60 * 60
    elif str_time[1].lower() == 'd':
        return int(str_time[0]) * 60 * 60 * 24
    else:
        return 0


def remove_figs():
    files = glob.glob('figures/*.png')
    for f in files:
        try:
            os.remove(f)
        except OSError as e:
            print("ERROR: %s : %s" % (f, e.strerror))


def report(unit):
    document = Document()
    document.add_heading('Report', 0)
    filename = "figures/" + str(unit.name) + "_*" + ".png"
    files = glob.glob(filename)
    for f in files:
        document.add_picture(f, width=Inches(6))
    document.save('report.docx')

# doctpl does not suport special characters like -, . ,etc
def fill_template(units):
    doc = DocxTemplate("template.docx")

    # # create temp docx to insert variables required for graphs
    context = {}
    for unit in units:
        context_name = "unit_" + str(unit.id) + "_" + str(unit.name)
        replacement = ""
        filename = "figures/" + str(unit.name) + "_*" + ".png"
        files = glob.glob(filename)
        for f in files:
            replacement += "{{ " + f[8:-4] + " }}\n"
        context[context_name] = replacement
    doc.render(context)
    doc.save("temp.docx")

    doc1 = DocxTemplate("temp.docx")
    figures = glob.glob("figures/*.png")
    context = {}
    for figure in figures:
        context[figure[8:-4]] = InlineImage(doc1, figure, width=Inches(6))
    doc1.render(context)
    doc1.save("generated_report.docx")
    # os.remove("temp.docx")

if "__main__" == __name__:
    # remove_figs()

    new_config = Config("template.docx")
    interval = get_seconds('1d')
    unit_id_name = new_config.get_units()
    unit_count = len(unit_id_name)

    units = [Unit(unit_id_name[i][0], unit_id_name[i][1]) for i in range(unit_count)]

    # for unit in units:
    #     unit.get_data(interval)
    #     unit.plot(['bvolt', 'bocv'], multiple_axes=True)
    #     unit.plot(['temp', 'InvTemp'])
    #     unit.plot(['pcvolt'])
    #     unit.plot(['bpow'])
    #     unit.plot(['prspps', 'prsnps'])
    #     unit.plot(['soc'])

    fill_template(units)
